"""
PO Bot — Two-Stage Pipeline
===========================
Stage 1 (MODE=standardise): Raw docs → clean product document → standardised/
Stage 2 (MODE=generate):    Clean doc → structured backlog JSON → backlog/
"""

import os, sys, json
from datetime import datetime, timezone

# ---------------------------------------------------------------------------
# TEXT EXTRACTION
# ---------------------------------------------------------------------------

def extract_text(filepath):
    ext = filepath.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        import PyPDF2
        with open(filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            text = "\n\n".join(p.extract_text() or "" for p in reader.pages).strip()
        if not text:
            raise ValueError("PDF appears image-based. Use a text-based PDF or .txt instead.")
        return text
    elif ext == "docx":
        from docx import Document
        doc = Document(filepath)
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        for t in doc.tables:
            for r in t.rows:
                lines.append(" | ".join(c.text.strip() for c in r.cells))
        return "\n".join(lines)
    elif ext in ("txt", "md"):
        return open(filepath, encoding="utf-8").read()
    else:
        raise ValueError(f"Unsupported format: .{ext}  (use .pdf .docx .md .txt)")


# ---------------------------------------------------------------------------
# STAGE 1 — STANDARDISE
# ---------------------------------------------------------------------------

STANDARDISE_SYSTEM = """You are a senior game product manager and technical producer at a Unity game studio.
Your job is to take ANY raw input — rough notes, Slack messages, voice memo transcripts, bullet points,
or partial specs — and convert it into a clean, complete game product document ready for sprint planning.

RULES:
- Extract and organise ALL information given — nothing is too small to capture
- Fill gaps intelligently using game development common sense
- If you make an assumption, mark it clearly: [Assumed: ...]
- Be SPECIFIC — "add enemies" is useless. "Add 3 enemy types: melee grunt, ranged archer, boss" is useful
- Think in terms of what each team needs to actually build this:
    Unity Dev: scenes, scripts, physics, animation controllers, UI, builds, store
    3D Team: models, textures, rigs, animations, VFX, audio, sound effects
    Backend: APIs, databases, cloud saves, leaderboards, analytics, auth
- Surface ALL technical constraints, platform targets, and performance goals
- Output ONLY the structured document — no preamble, no explanation"""

STANDARDISE_PROMPT = """Convert the following raw input into a clean, structured game product document.

Use EXACTLY this format:

# [Game Title]

## Game Overview
[2-3 sentences: genre, core loop, platform, and what makes it fun/unique]

## Problem / Opportunity
[Why does this game need to exist? What gap does it fill? Who is the target player?]

## Target Players
- **[Player Type 1]**: [age, gaming experience, what they want from this game]
- **[Player Type 2]**: [age, gaming experience, what they want from this game]

## Game Design

### Core Gameplay Loop
[Step-by-step description of the main gameplay cycle — what the player does every session]

### Game Mechanics
- [Specific mechanic with enough detail to implement — e.g. "Double-jump with 0.3s coyote time"]
- [Specific mechanic]

### Progression System
- [How the player advances — levels, XP, unlocks, difficulty curve]

### Win / Lose Conditions
- [Clear win state]
- [Clear fail state and consequences]

## Visual & Audio Direction

### Art Style
[Describe the look — reference games or styles, colour palette, perspective (2D/3D/isometric), tone]

### 3D / 2D Assets Required
- Characters: [list with description]
- Environments: [list with description]
- UI elements: [list with description]
- VFX: [list — particles, shaders, effects]

### Audio
- Music: [style, number of tracks, when they play]
- SFX: [key sound effects needed]

## Technical Specifications
- Engine: Unity [version if specified]
- Platform: [iOS / Android / PC / Console / WebGL]
- Target frame rate: [e.g. 60fps mobile, 120fps PC]
- Orientation: [Portrait / Landscape / Both]
- Min device spec: [e.g. iPhone X, Android 8.0]
- Multiplayer: [Yes/No — if yes, real-time or turn-based]
- Backend services: [cloud saves / leaderboards / analytics / auth / none]

## Monetisation (if applicable)
- Model: [Premium / Free-to-play / Ads / IAP]
- [Specific monetisation mechanic]

## Features by Area

### Core Gameplay
- [Specific implementable feature]

### UI / UX
- [Specific screen or flow]

### Backend & Services
- [Specific backend feature]

### Store & Deployment
- [Platform-specific requirement]

## Out of Scope (v1)
- [Feature explicitly excluded from this version]

## Open Questions
- [Anything unclear that needs a decision before work starts]

---
RAW INPUT:
{raw_text}
---"""


def stage1_standardise(raw_text, doc_name):
    import anthropic
    client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
    print("  Stage 1: Standardising document...")

    msg = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=3000,
        system=STANDARDISE_SYSTEM,
        messages=[{"role": "user", "content": STANDARDISE_PROMPT.format(raw_text=raw_text[:10000])}]
    )
    return msg.content[0].text.strip()


# ---------------------------------------------------------------------------
# STAGE 2 — GENERATE BACKLOG
# ---------------------------------------------------------------------------

BACKLOG_SYSTEM = """You are a Principal Product Owner embedded in a Unity game studio with 15 years of Agile/Scrum experience.
You understand exactly how game development works and how to write stories that engineers and artists can act on immediately.

PROCESS:
1. Read the full game document — understand the genre, loop, platforms, art style
2. Write a one-sentence vision that captures what makes this game worth building
3. Create 4-6 Epics — only ones directly supported by the document:
   - Core Gameplay Mechanics (physics, controls, game loop)
   - Characters & Animation (player, NPCs, enemies, rigs)
   - Environments & Level Design (scenes, layouts, props)
   - UI / UX (menus, HUD, screens, flows)
   - Audio & Visual FX (SFX, music, particles, shaders)
   - Backend & Services (cloud saves, leaderboards, auth, analytics)
   - Store & Release (build pipeline, platform submission, certificates)
   - QA & Polish (bug fixes, performance, playtesting)
4. Write 2-4 stories per epic — each story is ONE deliverable unit
5. Write exactly 2 Acceptance Criteria per story using Given/When/Then
6. Story points: 1=trivial 2=small 3=medium 5=large 8=complex 13=very complex
7. MoSCoW priorities: Must Have=MVP launch blocker, Should Have=important but not blocking, Could Have=nice-to-have, Wont Have=explicitly out of scope
8. Sprint planning:
   - Sprint 1: Walking skeleton — one playable scene, basic movement, placeholder art, project builds and runs on target platform
   - Sprint 2-3: All Must Have stories complete — core loop is fun and testable
   - Sprint 4+: Should Have and Could Have stories

RULES:
- Story titles must be action-oriented and specific: "Implement double-jump with coyote time" not "Add jumping"
- Acceptance criteria must be testable by a QA person or playtester
- Every story belongs to exactly ONE team — assign based on the actual implementation work
- Never write a story that two teams need to implement together — split it into two stories
- Base ONLY on the document content. Story IDs: US-001... Epic IDs: E1...
- OUTPUT: valid JSON only — no markdown, no explanation, start with {{ end with }}

TEAM ASSIGNMENT RULES (assign "team" field to every story):
- "unity_dev"  → C# scripting, MonoBehaviours, Unity scenes, Prefabs, Physics (Rigidbody/Colliders), Animation Controllers (Animator), UI Canvas/uGUI, TextMeshPro, NavMesh, Cinemachine, Input System, Build pipeline, iOS/Android deployment, Unity Store integration, any programming task
- "3d_team"    → Blender/Maya models, UV unwrapping, PBR textures, skeletal rigs, keyframe animations, blend shapes, particle systems, VFX Graph, Shader Graph materials, environmental art, character design, sound effects files, music composition, audio mixing
- "backend"    → REST APIs, Firebase/PlayFab/custom server, SQL/NoSQL databases, cloud save systems, leaderboard APIs, matchmaking, authentication, push notifications, analytics events, webhook handlers, server-side game logic"""

BACKLOG_PROMPT = """Return a complete sprint-ready product backlog as JSON.
Every field is mandatory. Every story must have a "team" field.

{{
  "project_name": "string — the game title",
  "vision": "string — one sentence: what game this is, for whom, and why it is compelling",
  "personas": ["string — player type with brief description"],
  "epics": [
    {{
      "id": "E1",
      "name": "string — epic name (e.g. Core Gameplay Mechanics)",
      "description": "string — what this epic covers and why it matters for the game",
      "stories": [
        {{
          "id": "US-001",
          "title": "string — specific action verb + object, max 8 words",
          "user_story": "As a [persona], I want [specific feature], so that [clear player benefit]",
          "acceptance_criteria": [
            "Given [context] When [player action] Then [specific measurable outcome]",
            "Given [context] When [edge case] Then [expected behaviour]"
          ],
          "story_points": 3,
          "priority": "Must Have",
          "sprint": 1,
          "team": "unity_dev",
          "notes": "string — technical notes, dependencies, or open questions (empty string if none)"
        }}
      ]
    }}
  ]
}}

GAME DOCUMENT:
---
{doc_text}
---"""


def stage2_generate(doc_text, doc_name):
    import anthropic
    client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
    print("  Stage 2: Generating backlog...")

    msg = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=16000,
        system=BACKLOG_SYSTEM,
        messages=[{"role": "user", "content": BACKLOG_PROMPT.format(doc_text=doc_text[:12000])}]
    )
    raw = msg.content[0].text.strip()
    if raw.startswith("```"):
        parts = raw.split("```")
        raw = parts[1][4:] if parts[1].startswith("json") else parts[1]
    return json.loads(raw.strip())


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------

def main():
    mode     = os.environ.get("MODE", "standardise").strip()
    doc_file = os.environ.get("DOC_FILE", "").strip()
    doc_name = os.environ.get("DOC_NAME", "doc").strip()
    extra    = os.environ.get("EXTRA_FILES", "").strip()

    if not doc_file or not os.path.exists(doc_file):
        print(f"DOC_FILE not set or not found: {doc_file}")
        sys.exit(0)

    ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    print(f"\n{'='*50}\n  PO Bot — {mode.upper()}\n  File: {doc_file}\n{'='*50}")

    print("\n[1] Extracting text...")
    raw_text = extract_text(doc_file)

    if extra:
        for extra_file in extra.split(","):
            extra_file = extra_file.strip()
            if extra_file and os.path.exists(extra_file):
                print(f"  + Merging: {extra_file}")
                raw_text += "\n\n---\n\n" + extract_text(extra_file)

    print(f"  {len(raw_text):,} characters total")

    if mode == "standardise":
        print("\n[2] Standardising with Claude...")
        clean_doc = stage1_standardise(raw_text, doc_name)

        os.makedirs("standardised", exist_ok=True)
        out = f"standardised/{doc_name}_{ts}.md"
        with open(out, "w", encoding="utf-8") as f:
            f.write(clean_doc)
        print(f"  Saved: {out}")
        print(f"\n{'='*50}\n  Stage 1 complete: {out}\n{'='*50}\n")

    elif mode == "generate":
        print("\n[2] Generating backlog with Claude...")
        backlog = stage2_generate(raw_text, doc_name)

        stories = sum(len(e["stories"]) for e in backlog["epics"])
        sprints = max(s["sprint"] for e in backlog["epics"] for s in e["stories"])
        print(f"  Project: {backlog.get('project_name')}")
        print(f"  Epics: {len(backlog['epics'])}  Stories: {stories}  Sprints: {sprints}")

        backlog["_meta"] = {
            "source_document": doc_file,
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "status": "pending_review"
        }

        os.makedirs("backlog", exist_ok=True)
        out = f"backlog/{doc_name}_{ts}.json"
        with open(out, "w", encoding="utf-8") as f:
            json.dump(backlog, f, indent=2, ensure_ascii=False)
        print(f"  Saved: {out}")

        print(f"\n  Sprint summary:")
        for n in range(1, sprints + 1):
            ss = [s for e in backlog["epics"] for s in e["stories"] if s["sprint"] == n]
            pts = sum(s["story_points"] for s in ss)
            must = sum(1 for s in ss if s["priority"] == "Must Have")
            print(f"    Sprint {n}: {len(ss)} stories ({pts} pts) - {must} Must Have")

        print(f"\n{'='*50}\n  Stage 2 complete: {out}\n{'='*50}\n")

    else:
        print(f"Unknown MODE: {mode}. Use 'standardise' or 'generate'.")
        sys.exit(1)


if __name__ == "__main__":
    main()
